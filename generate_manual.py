#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 JimuReport 使用操作说明文档
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_chinese_font(run, font_name='宋体', font_size=10.5, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(font_size)
    font.bold = bold
    if bold:
        font.color.rgb = RGBColor(0, 0, 0)

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    font_names = ['黑体', '黑体', '楷体', '宋体']
    font_sizes = [16, 14, 12, 11]
    set_chinese_font(run, font_names[level-1] if level <= 4 else '宋体', 
                     font_sizes[level-1] if level <= 4 else 11, bold=True)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_zh(doc, text, bold=False, font_size=10.5, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """添加中文段落"""
    para = doc.add_paragraph()
    para.alignment = alignment
    run = para.add_run(text)
    set_chinese_font(run, '宋体', font_size, bold)
    return para

def add_image_with_caption(doc, image_path, caption, max_width=6.0):
    """添加图片和说明"""
    if os.path.exists(image_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(max_width))
        
        # 添加图片说明
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_para.add_run(caption)
        set_chinese_font(run, '宋体', 9)
        
        doc.add_paragraph()  # 空行
        return True
    else:
        print(f"图片不存在: {image_path}")
        return False

def generate_manual():
    """生成使用操作说明文档"""
    doc = Document()
    
    # 设置默认中文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    
    # ===== 封面 =====
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('\n\n\n\n')
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('JimuReport 积木报表')
    set_chinese_font(run, '黑体', 26, bold=True)
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('使用操作说明')
    set_chinese_font(run, '黑体', 26, bold=True)
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run('\n\n数据源配置与 API 接口创建指南')
    set_chinese_font(run, '楷体', 16)
    
    doc.add_page_break()
    
    # ===== 目录 =====
    add_heading_zh(doc, '目录', level=1)
    
    toc_items = [
        '一、系统登录',
        '二、数据源管理',
        '    2.1 进入数据源列表',
        '    2.2 新增数据源',
        '    2.3 配置数据源信息',
        '    2.4 测试连接',
        '三、API 工作台',
        '    3.1 进入 API 生成器',
        '    3.2 创建 API 接口',
        '    3.3 配置 API 基础信息',
        '    3.4 选择数据源和表',
        '    3.5 配置字段和参数',
        '    3.6 测试 API 接口',
        '四、注意事项'
    ]
    
    for item in toc_items:
        para = doc.add_paragraph(item, style='List Number' if not item.startswith('    ') else 'List Bullet')
        for run in para.runs:
            set_chinese_font(run, '宋体', 11)
    
    doc.add_page_break()
    
    # ===== 正文开始 =====
    
    # 一、系统登录
    add_heading_zh(doc, '一、系统登录', level=1)
    
    add_paragraph_zh(doc, '1. 打开浏览器，访问 JimuReport 报表平台地址：')
    add_paragraph_zh(doc, '   http://localhost:8080/jmreport/list', bold=True)
    add_paragraph_zh(doc, '2. 在登录页面输入账号密码：')
    add_paragraph_zh(doc, '   • 用户名：admin')
    add_paragraph_zh(doc, '   • 密码：123456')
    add_paragraph_zh(doc, '3. 点击"登录"按钮进入系统')
    
    add_image_with_caption(doc, 'extracted_images/image1.png', '图 1-1 登录页面')
    
    # 二、数据源管理
    add_heading_zh(doc, '二、数据源管理', level=1)
    
    add_paragraph_zh(doc, '数据源是报表和 API 接口的数据基础，在使用前需要先配置数据源连接。')
    
    # 2.1 进入数据源列表
    add_heading_zh(doc, '2.1 进入数据源列表', level=2)
    add_paragraph_zh(doc, '登录系统后，在左侧菜单中找到"数据源"或"数据集"菜单，点击进入数据源管理页面。在数据源列表页面可以查看已配置的数据源，也可以进行新增、编辑、删除等操作。')
    
    add_image_with_caption(doc, 'extracted_images/image2.png', '图 2-1 数据源列表页面')
    
    # 2.2 新增数据源
    add_heading_zh(doc, '2.2 新增数据源', level=2)
    add_paragraph_zh(doc, '1. 在数据源列表页面，点击"新增"按钮')
    add_paragraph_zh(doc, '2. 在弹出的选项中选择"SQL 数据集"（用于数据库连接）')
    
    add_image_with_caption(doc, 'extracted_images/image3.png', '图 2-2 点击新增按钮')
    add_image_with_caption(doc, 'extracted_images/image4.png', '图 2-3 选择 SQL 数据集')
    
    # 2.3 配置数据源信息
    add_heading_zh(doc, '2.3 配置数据源信息', level=2)
    add_paragraph_zh(doc, '在数据源配置页面，需要填写以下信息：')
    
    add_paragraph_zh(doc, '【基础信息】')
    add_paragraph_zh(doc, '• 数据源名称：自定义名称，如"测试数据库"')
    add_paragraph_zh(doc, '• 数据源编码：唯一标识，如"test_db"')
    
    add_paragraph_zh(doc, '【数据库连接信息】')
    add_paragraph_zh(doc, '• 数据库类型：选择 MySQL、Oracle、SQL Server 等')
    add_paragraph_zh(doc, '• 驱动类：根据数据库类型自动填充或手动填写')
    add_paragraph_zh(doc, '• 数据库 URL：JDBC 连接地址')
    add_paragraph_zh(doc, '• 用户名：数据库账号')
    add_paragraph_zh(doc, '• 密码：数据库密码')
    
    add_image_with_caption(doc, 'extracted_images/image5.png', '图 2-4 维护数据源')
    add_image_with_caption(doc, 'extracted_images/image6.png', '图 2-5 新增数据源表单')
    
    # 2.4 测试连接
    add_heading_zh(doc, '2.4 测试连接', level=2)
    add_paragraph_zh(doc, '1. 填写完数据源信息后，点击"测试连接"按钮')
    add_paragraph_zh(doc, '2. 系统会尝试连接数据库，如果连接成功会显示提示信息')
    add_paragraph_zh(doc, '3. 测试成功后，点击"保存"按钮完成数据源创建')
    
    add_image_with_caption(doc, 'extracted_images/image7.png', '图 2-6 连接测试成功提示')
    
    # 三、API 工作台
    add_heading_zh(doc, '三、API 工作台', level=1)
    
    add_paragraph_zh(doc, 'API 工作台（API 生成器）是 JimuReport 的重要功能，可以通过可视化配置快速生成 REST API 接口，无需编写代码即可对外提供数据服务。')
    
    # 3.1 进入 API 生成器
    add_heading_zh(doc, '3.1 进入 API 生成器', level=2)
    add_paragraph_zh(doc, '1. 在系统菜单中找到"API 工作台"或"API 生成器"')
    add_paragraph_zh(doc, '2. 点击进入 API 生成器列表页面')
    add_paragraph_zh(doc, '3. 页面会显示已创建的 API 接口列表')
    
    add_image_with_caption(doc, 'extracted_images/image8.png', '图 3-1 API 生成器列表页')
    
    # 3.2 创建 API 接口
    add_heading_zh(doc, '3.2 创建 API 接口', level=2)
    add_paragraph_zh(doc, '在 API 生成器列表页面，点击"创建 API"或"新建"按钮，开始创建新的 API 接口。')
    
    # 3.3 配置 API 基础信息
    add_heading_zh(doc, '3.3 配置 API 基础信息', level=2)
    add_paragraph_zh(doc, '在创建 API 页面，首先需要配置基础信息：')
    
    add_paragraph_zh(doc, '【基础配置】')
    add_paragraph_zh(doc, '• API 名称：接口的名称，如"用户列表接口"')
    add_paragraph_zh(doc, '• API 路径：接口的访问路径，如"/api/user/list"')
    add_paragraph_zh(doc, '• API 类型：选择"表类型"或"SQL 类型"')
    add_paragraph_zh(doc, '• 描述：接口的功能说明')
    
    add_paragraph_zh(doc, '【分页配置】')
    add_paragraph_zh(doc, '• 是否分页：开启后接口支持分页查询')
    add_paragraph_zh(doc, '• 每页条数：默认每页返回的数据条数')
    
    add_image_with_caption(doc, 'extracted_images/image9.png', '图 3-2 创建 API - 基础信息配置')
    
    # 3.4 选择数据源和表
    add_heading_zh(doc, '3.4 选择数据源和表', level=2)
    add_paragraph_zh(doc, '1. 在数据源下拉框中选择之前创建的数据源')
    add_paragraph_zh(doc, '2. 选择要操作的数据库表')
    add_paragraph_zh(doc, '3. 如果是多表关联，可以配置表之间的关联关系（JOIN）')
    
    add_paragraph_zh(doc, '【表配置选项】')
    add_paragraph_zh(doc, '• 数据库名称：数据源中的具体数据库')
    add_paragraph_zh(doc, '• 表名：选择要查询的主表')
    add_paragraph_zh(doc, '• 表别名：为表设置别名，便于多表关联时使用')
    add_paragraph_zh(doc, '• 关联类型：如 INNER JOIN、LEFT JOIN 等')
    add_paragraph_zh(doc, '• 关联条件：设置表与表之间的关联字段')
    
    add_image_with_caption(doc, 'extracted_images/image10.png', '图 3-3 选择数据源和表配置')
    
    # 3.5 配置字段和参数
    add_heading_zh(doc, '3.5 配置字段和参数', level=2)
    
    add_paragraph_zh(doc, '【返回字段配置】')
    add_paragraph_zh(doc, '1. 在字段列表中勾选需要返回的字段')
    add_paragraph_zh(doc, '2. 可以为字段设置别名，如将"create_time"显示为"createTime"')
    add_paragraph_zh(doc, '3. 设置字段的数据类型和描述信息')
    
    add_paragraph_zh(doc, '【请求参数配置】')
    add_paragraph_zh(doc, '1. 点击"添加参数"按钮')
    add_paragraph_zh(doc, '2. 配置参数信息：')
    add_paragraph_zh(doc, '   • 参数名：如"userName"、"status"')
    add_paragraph_zh(doc, '   • 参数类型：query（查询参数）、path（路径参数）、body（请求体）')
    add_paragraph_zh(doc, '   • 数据类型：String、Integer、Date 等')
    add_paragraph_zh(doc, '   • 默认值：参数的默认值')
    add_paragraph_zh(doc, '   • 描述：参数的说明')
    add_paragraph_zh(doc, '   • 校验规则：如必填、长度限制等')
    
    add_image_with_caption(doc, 'extracted_images/image11.png', '图 3-4 字段配置和参数配置')
    
    add_paragraph_zh(doc, '3. 配置完成后，点击"保存"按钮保存 API 配置')
    add_paragraph_zh(doc, '4. 点击"发布"或"启用"按钮使接口生效')
    
    # 3.6 测试 API 接口
    add_heading_zh(doc, '3.6 测试 API 接口', level=2)
    add_paragraph_zh(doc, '1. 在 API 列表中找到刚创建的接口')
    add_paragraph_zh(doc, '2. 点击"测试"或"调试"按钮进入测试页面')
    add_paragraph_zh(doc, '3. 在测试页面填写请求参数值')
    add_paragraph_zh(doc, '4. 点击"发送请求"按钮')
    add_paragraph_zh(doc, '5. 查看返回结果，验证接口是否正确')
    
    add_image_with_caption(doc, 'extracted_images/image12.png', '图 3-5 API 测试页面')
    
    add_paragraph_zh(doc, '【返回结果说明】')
    add_paragraph_zh(doc, '接口返回标准的 JSON 格式数据，包含：')
    add_paragraph_zh(doc, '• code：状态码，200 表示成功')
    add_paragraph_zh(doc, '• message：提示信息')
    add_paragraph_zh(doc, '• data：返回的数据内容')
    add_paragraph_zh(doc, '• total：总记录数（分页时返回）')
    
    # 四、注意事项
    add_heading_zh(doc, '四、注意事项', level=1)
    
    add_paragraph_zh(doc, '1. 【数据源安全】数据库密码等敏感信息会加密存储，请妥善保管系统访问权限')
    add_paragraph_zh(doc, '2. 【API 路径规范】API 路径建议以"/api/"开头，避免与系统接口冲突')
    add_paragraph_zh(doc, '3. 【SQL 注入防护】系统会对 SQL 类型 API 进行安全检查，请勿尝试危险操作')
    add_paragraph_zh(doc, '4. 【性能优化】大数据量查询建议开启分页，避免一次性返回过多数据')
    add_paragraph_zh(doc, '5. 【权限控制】API 接口默认需要登录后才能访问，可在系统中配置权限策略')
    add_paragraph_zh(doc, '6. 【版本管理】修改已发布的 API 时请谨慎，可能影响正在使用的客户端')
    add_paragraph_zh(doc, '7. 【日志记录】系统会记录 API 调用日志，便于问题排查和审计')
    
    # 保存文档
    output_path = 'JimuReport使用操作说明.docx'
    doc.save(output_path)
    print(f"文档已生成: {os.path.abspath(output_path)}")
    return output_path

if __name__ == "__main__":
    generate_manual()
